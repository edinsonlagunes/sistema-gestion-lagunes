import io

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from app import models, schemas
from app.auth import obtener_usuario_actual, requerir_admin
from app.database import get_db
from app.zona_horaria import ahora_peru

router = APIRouter(prefix="/insumos", tags=["Insumos"], dependencies=[Depends(obtener_usuario_actual)])
proveedores_router = APIRouter(
    prefix="/proveedores", tags=["Proveedores"], dependencies=[Depends(obtener_usuario_actual)]
)


@router.get("/", response_model=list[schemas.Insumo])
def listar_insumos(negocio_id: int | None = None, db: Session = Depends(get_db)):
    query = db.query(models.Insumo)
    if negocio_id is not None:
        query = query.filter(models.Insumo.negocio_id == negocio_id)
    return query.all()


@router.post("/", response_model=schemas.Insumo)
def crear_insumo(data: schemas.InsumoCreate, db: Session = Depends(get_db)):
    if not db.query(models.Negocio).get(data.negocio_id):
        raise HTTPException(status_code=404, detail="Negocio no encontrado")
    insumo = models.Insumo(**data.model_dump())
    db.add(insumo)
    db.commit()
    db.refresh(insumo)
    return insumo


@proveedores_router.get("/", response_model=list[schemas.Proveedor])
def listar_proveedores(db: Session = Depends(get_db)):
    return db.query(models.Proveedor).all()


@proveedores_router.post("/", response_model=schemas.Proveedor)
def crear_proveedor(data: schemas.ProveedorCreate, db: Session = Depends(get_db)):
    proveedor = models.Proveedor(**data.model_dump())
    db.add(proveedor)
    db.commit()
    db.refresh(proveedor)
    return proveedor


@proveedores_router.patch("/{proveedor_id}", response_model=schemas.Proveedor)
def editar_proveedor(
    proveedor_id: int,
    data: schemas.ProveedorUpdate,
    db: Session = Depends(get_db),
    _admin: models.Usuario = Depends(requerir_admin),
):
    """Edita los datos de un proveedor. Solo administradores."""
    proveedor = db.query(models.Proveedor).get(proveedor_id)
    if not proveedor:
        raise HTTPException(status_code=404, detail="Proveedor no encontrado")
    for campo, valor in data.model_dump(exclude_unset=True).items():
        setattr(proveedor, campo, valor)
    db.commit()
    db.refresh(proveedor)
    return proveedor


@proveedores_router.delete("/{proveedor_id}", status_code=204)
def eliminar_proveedor(
    proveedor_id: int,
    db: Session = Depends(get_db),
    _admin: models.Usuario = Depends(requerir_admin),
):
    """
    Quita un proveedor del catálogo. Bloqueado si tiene compras o pagos
    registrados — esos son historial financiero real y no se pueden
    perder solo porque se quiere quitar el proveedor de la lista.
    Solo administradores.
    """
    proveedor = db.query(models.Proveedor).get(proveedor_id)
    if not proveedor:
        raise HTTPException(status_code=404, detail="Proveedor no encontrado")
    if proveedor.compras or proveedor.pagos:
        raise HTTPException(
            status_code=400,
            detail="Este proveedor tiene compras o pagos registrados y no se puede eliminar. "
            "Si ya no trabajas con él, puedes dejarlo en la lista sin usarlo.",
        )
    db.delete(proveedor)
    db.commit()
    return None


@proveedores_router.post("/importar", response_model=schemas.ProveedorImportarResultado)
def importar_proveedores(
    items: list[schemas.ProveedorImportarItem],
    db: Session = Depends(get_db),
    _admin: models.Usuario = Depends(requerir_admin),
):
    """
    Crea proveedores en bloque (desde un Excel/CSV ya parseado en el
    frontend). Omite los que coincidan por nombre exacto con uno que ya
    existe, para no duplicar. Solo administradores.
    """
    existentes = {p.nombre.strip().lower() for p in db.query(models.Proveedor).all()}
    creados = []
    omitidos = 0
    for item in items:
        nombre_limpio = item.nombre.strip()
        if not nombre_limpio or nombre_limpio.lower() in existentes:
            omitidos += 1
            continue
        proveedor = models.Proveedor(
            nombre=nombre_limpio,
            contacto=item.contacto or None,
            telefono=item.telefono or None,
            direccion=item.direccion or None,
        )
        db.add(proveedor)
        existentes.add(nombre_limpio.lower())
        creados.append(proveedor)
    db.commit()
    for p in creados:
        db.refresh(p)
    return schemas.ProveedorImportarResultado(creados=len(creados), omitidos=omitidos, proveedores=creados)


@proveedores_router.get("/exportar-word")
def exportar_proveedores_word(db: Session = Depends(get_db)):
    """
    Genera un documento Word con la lista de proveedores (nombre,
    teléfono, dirección, contacto) — para imprimir o compartir fuera
    del sistema. El Excel y el PDF se generan directo en el navegador;
    este es el único que necesita armarse en el servidor.
    """
    from docx import Document
    from docx.shared import Pt

    proveedores = db.query(models.Proveedor).order_by(models.Proveedor.nombre).all()

    documento = Document()
    documento.add_heading("Proveedores", level=1)
    documento.add_paragraph(f"Generado el {ahora_peru().strftime('%d/%m/%Y %H:%M')}")

    tabla = documento.add_table(rows=1, cols=4)
    tabla.style = "Light Grid Accent 1"
    encabezados = tabla.rows[0].cells
    for celda, texto in zip(encabezados, ["Nombre", "Teléfono", "Dirección", "Contacto"]):
        celda.text = texto
        celda.paragraphs[0].runs[0].font.bold = True

    for p in proveedores:
        fila = tabla.add_row().cells
        fila[0].text = p.nombre or ""
        fila[1].text = p.telefono or ""
        fila[2].text = p.direccion or ""
        fila[3].text = p.contacto or ""

    for fila in tabla.rows:
        for celda in fila.cells:
            for parrafo in celda.paragraphs:
                for run in parrafo.runs:
                    run.font.size = Pt(10)

    buffer = io.BytesIO()
    documento.save(buffer)
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=proveedores.docx"},
    )


def _totales_proveedor(proveedor: models.Proveedor, negocio_id: int | None = None) -> tuple[float, float]:
    compras = proveedor.compras
    pagos = proveedor.pagos
    if negocio_id is not None:
        compras = [c for c in compras if c.negocio_id == negocio_id]
        pagos = [p for p in pagos if p.negocio_id == negocio_id]
    total_comprado = sum(c.costo for c in compras)
    total_pagado = sum(p.monto for p in pagos)
    return total_comprado, total_pagado


@proveedores_router.get("/resumen-pagos", response_model=list[schemas.ResumenPagoProveedor])
def resumen_pagos_proveedores(negocio_id: int | None = None, db: Session = Depends(get_db)):
    """
    Para el Dashboard: cuentas por pagar — por cada proveedor, cuánto se
    le ha comprado, cuánto se le ha pagado, y cuánto falta. Solo incluye
    proveedores con saldo pendiente.
    """
    resultado = []
    for proveedor in db.query(models.Proveedor).all():
        total_comprado, total_pagado = _totales_proveedor(proveedor, negocio_id)
        saldo = total_comprado - total_pagado
        if saldo <= 0:
            continue

        pagos = proveedor.pagos if negocio_id is None else [p for p in proveedor.pagos if p.negocio_id == negocio_id]
        ultimo_pago = max(pagos, key=lambda p: p.fecha_pago) if pagos else None

        resultado.append(
            schemas.ResumenPagoProveedor(
                proveedor_id=proveedor.id,
                proveedor_nombre=proveedor.nombre,
                total_comprado=total_comprado,
                total_pagado=total_pagado,
                saldo_pendiente=saldo,
                ultimo_pago_monto=ultimo_pago.monto if ultimo_pago else None,
                ultimo_pago_fecha=ultimo_pago.fecha_pago if ultimo_pago else None,
            )
        )

    resultado.sort(key=lambda r: r.saldo_pendiente, reverse=True)
    return resultado


@proveedores_router.get("/{proveedor_id}", response_model=schemas.ProveedorDetalle)
def obtener_proveedor(proveedor_id: int, db: Session = Depends(get_db)):
    proveedor = db.query(models.Proveedor).get(proveedor_id)
    if not proveedor:
        raise HTTPException(status_code=404, detail="Proveedor no encontrado")
    total_comprado, total_pagado = _totales_proveedor(proveedor)
    return schemas.ProveedorDetalle(
        id=proveedor.id,
        nombre=proveedor.nombre,
        contacto=proveedor.contacto,
        total_comprado=total_comprado,
        total_pagado=total_pagado,
        saldo_pendiente=total_comprado - total_pagado,
        pagos=proveedor.pagos,
    )


@proveedores_router.post("/{proveedor_id}/pagos", response_model=schemas.ProveedorDetalle)
def registrar_pago_proveedor(
    proveedor_id: int,
    data: schemas.PagoProveedorCreate,
    negocio_id: int,
    db: Session = Depends(get_db),
    _admin: models.Usuario = Depends(requerir_admin),
):
    """
    Registra un pago real a un proveedor — genera el egreso en finanzas
    por el monto efectivamente pagado (a diferencia de comprar, que no
    mueve caja todavía). Solo administradores.
    """
    proveedor = db.query(models.Proveedor).get(proveedor_id)
    if not proveedor:
        raise HTTPException(status_code=404, detail="Proveedor no encontrado")
    if not db.query(models.Negocio).get(negocio_id):
        raise HTTPException(status_code=404, detail="Negocio no encontrado")

    pago = models.PagoProveedor(
        proveedor_id=proveedor.id,
        negocio_id=negocio_id,
        monto=data.monto,
        fecha_pago=data.fecha_pago or ahora_peru(),
        medio_pago=data.medio_pago,
        descripcion=data.descripcion,
    )
    db.add(pago)
    db.flush()

    db.add(
        models.Egreso(
            negocio_id=negocio_id,
            categoria="pago_proveedor",
            monto=pago.monto,
            descripcion=f"Pago a proveedor - {proveedor.nombre}" + (f" ({data.descripcion})" if data.descripcion else ""),
            fecha=pago.fecha_pago,
            pago_proveedor_id=pago.id,
        )
    )

    db.commit()
    db.refresh(proveedor)
    total_comprado, total_pagado = _totales_proveedor(proveedor)
    return schemas.ProveedorDetalle(
        id=proveedor.id,
        nombre=proveedor.nombre,
        contacto=proveedor.contacto,
        total_comprado=total_comprado,
        total_pagado=total_pagado,
        saldo_pendiente=total_comprado - total_pagado,
        pagos=proveedor.pagos,
    )


@proveedores_router.patch("/{proveedor_id}/pagos/{pago_id}", response_model=schemas.ProveedorDetalle)
def actualizar_pago_proveedor(
    proveedor_id: int,
    pago_id: int,
    data: schemas.PagoProveedorUpdate,
    db: Session = Depends(get_db),
    _admin: models.Usuario = Depends(requerir_admin),
):
    """Corrige un pago a proveedor ya registrado y su egreso vinculado. Solo administradores."""
    pago = (
        db.query(models.PagoProveedor)
        .filter(models.PagoProveedor.id == pago_id, models.PagoProveedor.proveedor_id == proveedor_id)
        .first()
    )
    if not pago:
        raise HTTPException(status_code=404, detail="Pago no encontrado")

    for campo, valor in data.model_dump(exclude_unset=True).items():
        setattr(pago, campo, valor)

    egreso = db.query(models.Egreso).filter(models.Egreso.pago_proveedor_id == pago.id).first()
    if egreso:
        egreso.monto = pago.monto
        egreso.fecha = pago.fecha_pago

    db.commit()
    proveedor = db.query(models.Proveedor).get(proveedor_id)
    total_comprado, total_pagado = _totales_proveedor(proveedor)
    return schemas.ProveedorDetalle(
        id=proveedor.id,
        nombre=proveedor.nombre,
        contacto=proveedor.contacto,
        total_comprado=total_comprado,
        total_pagado=total_pagado,
        saldo_pendiente=total_comprado - total_pagado,
        pagos=proveedor.pagos,
    )


@proveedores_router.delete("/{proveedor_id}/pagos/{pago_id}", response_model=schemas.ProveedorDetalle)
def eliminar_pago_proveedor(
    proveedor_id: int,
    pago_id: int,
    db: Session = Depends(get_db),
    _admin: models.Usuario = Depends(requerir_admin),
):
    """Quita un pago a proveedor registrado por error, junto con su egreso. Solo administradores."""
    pago = (
        db.query(models.PagoProveedor)
        .filter(models.PagoProveedor.id == pago_id, models.PagoProveedor.proveedor_id == proveedor_id)
        .first()
    )
    if not pago:
        raise HTTPException(status_code=404, detail="Pago no encontrado")

    db.query(models.Egreso).filter(models.Egreso.pago_proveedor_id == pago.id).delete()
    db.delete(pago)
    db.commit()
    proveedor = db.query(models.Proveedor).get(proveedor_id)
    total_comprado, total_pagado = _totales_proveedor(proveedor)
    return schemas.ProveedorDetalle(
        id=proveedor.id,
        nombre=proveedor.nombre,
        contacto=proveedor.contacto,
        total_comprado=total_comprado,
        total_pagado=total_pagado,
        saldo_pendiente=total_comprado - total_pagado,
        pagos=proveedor.pagos,
    )
